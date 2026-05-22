import io
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.pdfmetrics import registerFont


def _parse_markdown(content: str):
    """将 Markdown 文本解析为结构化段落列表。"""
    lines = content.splitlines()
    paragraphs = []
    in_table = False
    table_lines = []

    for line in lines:
        stripped = line.strip()

        # 表格检测
        if stripped.startswith("|") and stripped.endswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            continue
        elif in_table:
            in_table = False
            paragraphs.append({"type": "table", "lines": table_lines})
            table_lines = []

        if not stripped:
            paragraphs.append({"type": "blank"})
            continue

        # 标题
        if stripped.startswith("# "):
            paragraphs.append({"type": "h1", "text": stripped[2:].strip()})
        elif stripped.startswith("## "):
            paragraphs.append({"type": "h2", "text": stripped[3:].strip()})
        elif stripped.startswith("### "):
            paragraphs.append({"type": "h3", "text": stripped[4:].strip()})
        elif stripped.startswith("> "):
            paragraphs.append({"type": "quote", "text": stripped[2:].strip()})
        elif stripped.startswith("- "):
            paragraphs.append({"type": "bullet", "text": stripped[2:].strip()})
        elif re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            paragraphs.append({"type": "numbered", "text": text})
        elif stripped.startswith("---"):
            paragraphs.append({"type": "hr"})
        else:
            paragraphs.append({"type": "paragraph", "text": stripped})

    if in_table and table_lines:
        paragraphs.append({"type": "table", "lines": table_lines})

    return paragraphs


def _strip_md_inline(text: str) -> str:
    """去除简单的 Markdown 行内标记（**、*、`"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def export_to_docx(
    content: str,
    report_id: int,
    exam_name: Optional[str] = None,
    class_name: Optional[str] = None,
    created_at: Optional[datetime] = None,
) -> io.BytesIO:
    """将 Markdown 报告转换为 Word 文档并返回 BytesIO。"""
    doc = Document()

    # 设置默认中文字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 标题页
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI 学情分析报告")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph()

    if exam_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"考试：{exam_name}")
        run.font.size = Pt(14)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    if class_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"班级：{class_name}")
        run.font.size = Pt(14)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    if created_at:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"生成时间：{created_at.strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"报告编号：{report_id}")
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    paragraphs = _parse_markdown(content)

    for para in paragraphs:
        ptype = para.get("type")
        if ptype == "blank":
            continue
        elif ptype == "h1":
            p = doc.add_heading(level=1)
            run = p.add_run(_strip_md_inline(para["text"]))
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        elif ptype == "h2":
            p = doc.add_heading(level=2)
            run = p.add_run(_strip_md_inline(para["text"]))
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        elif ptype == "h3":
            p = doc.add_heading(level=3)
            run = p.add_run(_strip_md_inline(para["text"]))
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        elif ptype == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(_strip_md_inline(para["text"]))
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif ptype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_strip_md_inline(para["text"]))
        elif ptype == "numbered":
            p = doc.add_paragraph(style="List Number")
            p.add_run(_strip_md_inline(para["text"]))
        elif ptype == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        elif ptype == "table":
            _add_docx_table(doc, para["lines"])
        else:
            p = doc.add_paragraph()
            p.add_run(_strip_md_inline(para["text"]))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_docx_table(doc, lines):
    """将 Markdown 表格行添加到 Word 文档。"""
    if not lines:
        return
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    # 过滤分隔行
    rows = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r)]
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    for i, row_cells in enumerate(rows):
        for j in range(num_cols):
            cell = table.cell(i, j)
            text = row_cells[j] if j < len(row_cells) else ""
            cell.text = _strip_md_inline(text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def _get_pdf_chinese_style():
    """注册中文字体并返回样式表。"""
    # 尝试注册系统常见中文字体，否则使用 ReportLab 内置 CID 字体
    font_name = "STSong-Light"
    try:
        # 尝试注册 Windows 常见中文字体
        pdfmetrics.registerFont(TTFont("SimSun", "C:/Windows/Fonts/simsun.ttc", subfontIndex=0))
        font_name = "SimSun"
    except Exception:
        try:
            pdfmetrics.registerFont(TTFont("SimSun", "C:/Windows/Fonts/simsunb.ttf"))
            font_name = "SimSun"
        except Exception:
            try:
                pdfmetrics.registerFont(TTFont("MicrosoftYaHei", "C:/Windows/Fonts/msyh.ttc", subfontIndex=0))
                font_name = "MicrosoftYaHei"
            except Exception:
                try:
                    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
                    font_name = "STSong-Light"
                except Exception:
                    pass

    styles = getSampleStyleSheet()

    base_style = ParagraphStyle(
        "ChineseBase",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=18,
        wordWrap="CJK",
    )

    styles.add(base_style)
    styles.add(ParagraphStyle(
        "ChineseTitle",
        parent=base_style,
        fontSize=20,
        leading=28,
        alignment=1,
        spaceAfter=20,
        fontName=font_name,
    ))
    styles.add(ParagraphStyle(
        "ChineseH1",
        parent=base_style,
        fontSize=16,
        leading=24,
        spaceAfter=10,
        spaceBefore=14,
        fontName=font_name,
    ))
    styles.add(ParagraphStyle(
        "ChineseH2",
        parent=base_style,
        fontSize=14,
        leading=20,
        spaceAfter=8,
        spaceBefore=10,
        fontName=font_name,
    ))
    styles.add(ParagraphStyle(
        "ChineseH3",
        parent=base_style,
        fontSize=12,
        leading=18,
        spaceAfter=6,
        spaceBefore=8,
        fontName=font_name,
    ))
    styles.add(ParagraphStyle(
        "ChineseQuote",
        parent=base_style,
        fontSize=10,
        leading=16,
        leftIndent=20,
        textColor=colors.grey,
        fontName=font_name,
    ))
    styles.add(ParagraphStyle(
        "ChineseBullet",
        parent=base_style,
        fontSize=11,
        leading=18,
        leftIndent=20,
        bulletIndent=10,
        bulletFontName=font_name,
        fontName=font_name,
    ))
    styles.add(ParagraphStyle(
        "ChineseBody",
        parent=base_style,
        fontSize=11,
        leading=18,
        fontName=font_name,
    ))

    return styles, font_name


def _escape_xml(text: str) -> str:
    """转义 XML 特殊字符。"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _md_to_pdf_text(text: str, font_name: str) -> str:
    """将简单 Markdown 行内标记转换为 ReportLab 可识别的 <b>/<i> 标签。"""
    text = _escape_xml(text)
    # **bold**
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # *italic*
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    # `code`
    text = re.sub(r"`(.+?)`", r"\1", text)
    return f"<font name='{font_name}'>{text}</font>"


def export_to_pdf(
    content: str,
    report_id: int,
    exam_name: Optional[str] = None,
    class_name: Optional[str] = None,
    created_at: Optional[datetime] = None,
) -> io.BytesIO:
    """将 Markdown 报告转换为 PDF 并返回 BytesIO。"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )
    styles, font_name = _get_pdf_chinese_style()

    story = []

    # 标题页
    story.append(Paragraph("AI 学情分析报告", styles["ChineseTitle"]))
    story.append(Spacer(1, 0.5 * cm))

    if exam_name:
        story.append(Paragraph(f"考试：{exam_name}", styles["ChineseBody"]))
    if class_name:
        story.append(Paragraph(f"班级：{class_name}", styles["ChineseBody"]))
    if created_at:
        story.append(Paragraph(f"生成时间：{created_at.strftime('%Y-%m-%d %H:%M')}", styles["ChineseBody"]))
    story.append(Paragraph(f"报告编号：{report_id}", styles["ChineseBody"]))
    story.append(Spacer(1, 1 * cm))

    paragraphs = _parse_markdown(content)

    for para in paragraphs:
        ptype = para.get("type")
        if ptype == "blank":
            continue
        elif ptype == "h1":
            story.append(Paragraph(_md_to_pdf_text(para["text"], font_name), styles["ChineseH1"]))
        elif ptype == "h2":
            story.append(Paragraph(_md_to_pdf_text(para["text"], font_name), styles["ChineseH2"]))
        elif ptype == "h3":
            story.append(Paragraph(_md_to_pdf_text(para["text"], font_name), styles["ChineseH3"]))
        elif ptype == "quote":
            story.append(Paragraph(_md_to_pdf_text(para["text"], font_name), styles["ChineseQuote"]))
        elif ptype == "bullet":
            story.append(Paragraph(f"• {_md_to_pdf_text(para['text'], font_name)}", styles["ChineseBullet"]))
        elif ptype == "numbered":
            story.append(Paragraph(f"{_md_to_pdf_text(para['text'], font_name)}", styles["ChineseBullet"]))
        elif ptype == "hr":
            story.append(Spacer(1, 0.3 * cm))
            story.append(Table([[""]], colWidths=[15 * cm], rowHeights=[1], style=TableStyle([
                ("LINEBELOW", (0, 0), (-1, 0), 1, colors.grey),
            ])))
            story.append(Spacer(1, 0.3 * cm))
        elif ptype == "table":
            _add_pdf_table(story, para["lines"], font_name)
        else:
            story.append(Paragraph(_md_to_pdf_text(para["text"], font_name), styles["ChineseBody"]))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _add_pdf_table(story, lines, font_name):
    """将 Markdown 表格行添加到 PDF。"""
    if not lines:
        return
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    # 过滤分隔行
    rows = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r)]
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table_data = []
    for row_cells in rows:
        row = []
        for j in range(num_cols):
            text = row_cells[j] if j < len(row_cells) else ""
            row.append(Paragraph(_md_to_pdf_text(text, font_name), ParagraphStyle(
                "TableCell",
                fontName=font_name,
                fontSize=10,
                leading=14,
                wordWrap="CJK",
            )))
        table_data.append(row)

    col_width = 15 * cm / num_cols
    table = Table(table_data, colWidths=[col_width] * num_cols)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f2f2")),
    ]))
    story.append(Spacer(1, 0.2 * cm))
    story.append(table)
    story.append(Spacer(1, 0.2 * cm))
